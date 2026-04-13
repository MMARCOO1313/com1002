from __future__ import annotations

import json
import math
from datetime import datetime
from pathlib import Path
from textwrap import wrap

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(r"C:\Users\Administrator\Desktop\shatin project")
REPO = ROOT / "bridgespace-v2-clone"
OUTPUT = ROOT / "BridgeSpace_Final_Report_Updated.docx"
ASSET_DIR = REPO / "docs" / "report-assets"
ASSET_DIR.mkdir(parents=True, exist_ok=True)


def font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont:
    candidates = [
        ("C:/Windows/Fonts/segoeuib.ttf" if bold else "C:/Windows/Fonts/segoeui.ttf"),
        ("C:/Windows/Fonts/arialbd.ttf" if bold else "C:/Windows/Fonts/arial.ttf"),
    ]
    for candidate in candidates:
        path = Path(candidate)
        if path.exists():
            return ImageFont.truetype(str(path), size=size)
    return ImageFont.load_default()


def draw_wrapped(draw: ImageDraw.ImageDraw, text: str, xy: tuple[int, int], max_width: int, body_font, fill, line_gap=8):
    x, y = xy
    words = text.split()
    lines: list[str] = []
    current = ""
    for word in words:
        test = word if not current else f"{current} {word}"
        if draw.textlength(test, font=body_font) <= max_width:
            current = test
        else:
            if current:
                lines.append(current)
            current = word
    if current:
        lines.append(current)
    for line in lines:
        draw.text((x, y), line, font=body_font, fill=fill)
        y += body_font.size + line_gap
    return y


def make_architecture_diagram() -> Path:
    path = ASSET_DIR / "architecture-diagram.png"
    img = Image.new("RGB", (1800, 1040), "#0f172a")
    draw = ImageDraw.Draw(img)
    title_font = font(42, bold=True)
    box_title = font(28, bold=True)
    box_body = font(21)
    caption_font = font(18)
    orange = "#f97316"
    slate = "#1e293b"
    border = "#334155"
    white = "#f8fafc"
    cyan = "#38bdf8"
    green = "#22c55e"

    draw.text((70, 45), "BridgeSpace v2.0 Digital Architecture", font=title_font, fill=white)
    draw.text((70, 100), "Final implemented structure used in the verified demo build on 12 April 2026.", font=caption_font, fill="#cbd5e1")

    boxes = {
        "smartgate": (80, 210, 420, 420),
        "smartcount": (80, 500, 420, 710),
        "backend": (560, 250, 1240, 670),
        "dashboard": (1320, 220, 1700, 430),
        "control": (1320, 500, 1700, 710),
    }

    def card(key, title, bullets, accent):
        x1, y1, x2, y2 = boxes[key]
        draw.rounded_rectangle((x1, y1, x2, y2), radius=28, fill=slate, outline=accent, width=4)
        draw.text((x1 + 24, y1 + 24), title, font=box_title, fill=white)
        y = y1 + 78
        for bullet in bullets:
            draw.ellipse((x1 + 28, y + 10, x1 + 38, y + 20), fill=accent)
            y = draw_wrapped(draw, bullet, (x1 + 52, y), x2 - x1 - 78, box_body, white, 6) + 8

    card("smartgate", "SmartGate Kiosk", [
        "Local face-signature matching with registration, queue join, walk-in entry, and session extension.",
        "Calls /users/register, /queue/join, /session/enter, and /session/extend.",
    ], orange)
    card("smartcount", "SmartCount Vision", [
        "YOLO-based occupancy updates from webcam input.",
        "Feeds /zones/occupancy to drive queue advancement and status changes.",
    ], cyan)
    card("backend", "FastAPI Backend + SQLite", [
        "Central API and state store for users, zones, queue, sessions, device log, and alerts.",
        "SessionManager enforces timers, extension rules, warnings, expiry, and no-show handling.",
        "OccupancyWatcher decides when departures are stable enough to auto-call the next visitor.",
        "WebSocket broadcast keeps the dashboard live without manual refresh.",
    ], green)
    card("dashboard", "React Dashboard", [
        "Live occupancy, queue, session countdowns, device status, and alerts.",
        "Connected through /ws and REST fallback for initial zone load.",
    ], "#a855f7")
    card("control", "SmartControl + AlertEngine", [
        "Logs simulated device actions for gates, hoops, and lighting.",
        "Escalates incidents to dashboard, Telegram, or Twilio-ready alert channels.",
    ], "#eab308")

    def arrow(start, end, color):
        draw.line((start, end), fill=color, width=7)
        angle = math.atan2(end[1] - start[1], end[0] - start[0])
        arrow_size = 20
        p1 = (
            end[0] - arrow_size * math.cos(angle - math.pi / 6),
            end[1] - arrow_size * math.sin(angle - math.pi / 6),
        )
        p2 = (
            end[0] - arrow_size * math.cos(angle + math.pi / 6),
            end[1] - arrow_size * math.sin(angle + math.pi / 6),
        )
        draw.polygon([end, p1, p2], fill=color)

    arrow((420, 315), (560, 315), orange)
    arrow((420, 605), (560, 605), cyan)
    arrow((1240, 350), (1320, 350), "#a855f7")
    arrow((1240, 605), (1320, 605), "#eab308")
    arrow((1510, 430), (1510, 500), "#94a3b8")

    draw.text((650, 730), "Sense -> Decide -> Act -> Notify", font=box_title, fill="#fdba74")
    draw.text((650, 780), "The final project closes the loop automatically: occupancy changes trigger queue decisions, session timers trigger device and alert actions, and the dashboard reflects the result in real time.", font=caption_font, fill="#cbd5e1")

    img.save(path)
    return path


def make_operation_diagram() -> Path:
    path = ASSET_DIR / "operation-flow.png"
    img = Image.new("RGB", (1800, 900), "#fffaf3")
    draw = ImageDraw.Draw(img)
    title_font = font(42, bold=True)
    step_font = font(24, bold=True)
    body_font = font(19)
    draw.text((70, 45), "Verified User Operation Flow", font=title_font, fill="#111827")
    draw.text((70, 98), "Flow used during the final local verification on 12 April 2026.", font=body_font, fill="#475569")

    steps = [
        ("1. Arrival & Scan", "Visitor approaches the SmartGate kiosk. A returning user is recognised by face-signature match, while a new user can register name, phone, and face_id."),
        ("2. Zone Selection", "The kiosk sends the chosen zone to the backend. If occupancy is low and no queue exists, the backend returns walk_in = true."),
        ("3. Queue or Walk-in", "If a zone is busy, the user receives a queue number. If not, the kiosk allows immediate entry through /session/enter."),
        ("4. Live Monitoring", "SmartCount updates occupancy. The dashboard shows occupancy bars, queue state, session timers, device state, and alerts."),
        ("5. Timed Session", "SessionManager enforces zone-specific durations. The user can request one extension through the kiosk if queue conditions allow it."),
        ("6. Auto Progression", "When occupancy drops and departure is confirmed, OccupancyWatcher auto-calls the next visitor and pushes the event to the dashboard."),
    ]

    x_positions = [70, 340, 610, 880, 1150, 1420]
    y = 230
    colors = ["#ea580c", "#f97316", "#fb923c", "#f59e0b", "#14b8a6", "#0ea5e9"]
    for idx, ((title, text), x) in enumerate(zip(steps, x_positions)):
        draw.rounded_rectangle((x, y, x + 230, y + 420), radius=26, fill="#ffffff", outline=colors[idx], width=5)
        draw.rounded_rectangle((x + 16, y + 16, x + 78, y + 78), radius=16, fill=colors[idx], outline=colors[idx])
        draw.text((x + 34, y + 30), str(idx + 1), font=font(30, bold=True), fill="white")
        draw.text((x + 16, y + 104), title, font=step_font, fill="#111827")
        draw_wrapped(draw, text, (x + 16, y + 150), 198, body_font, "#334155", 7)
        if idx < len(steps) - 1:
            start = (x + 230, y + 210)
            end = (x + 255, y + 210)
            draw.line((start, end), fill="#94a3b8", width=7)
            draw.polygon([end, (end[0] - 18, end[1] - 12), (end[0] - 18, end[1] + 12)], fill="#94a3b8")

    draw.text((70, 720), "Practical outcome", font=step_font, fill="#111827")
    draw_wrapped(
        draw,
        "In the final verification run, the backend successfully handled registration, zone listing, occupancy updates, walk-in admission, session creation, and active session retrieval. This confirmed that the user-facing flow documented above is not only conceptual but executable in the current codebase.",
        (70, 760),
        1630,
        body_font,
        "#334155",
        8,
    )
    img.save(path)
    return path


def make_validation_panel() -> Path:
    path = ASSET_DIR / "validation-panel.png"
    img = Image.new("RGB", (1800, 1020), "#f8fafc")
    draw = ImageDraw.Draw(img)
    title_font = font(42, bold=True)
    section_font = font(28, bold=True)
    body_font = font(21)
    mono_font = font(19)
    draw.text((70, 45), "Final Practical Verification Snapshot", font=title_font, fill="#0f172a")
    draw.text((70, 100), "Evidence gathered during the final local run used for this report.", font=body_font, fill="#475569")

    left = (70, 170, 840, 920)
    right = (920, 170, 1730, 920)
    for rect in [left, right]:
        draw.rounded_rectangle(rect, radius=28, fill="white", outline="#cbd5e1", width=3)

    draw.text((left[0] + 28, left[1] + 24), "Startup and build checks", font=section_font, fill="#111827")
    checks = [
        "Backend served on 127.0.0.1:8012 and 127.0.0.1:8000.",
        "Frontend dev server served on 127.0.0.1:4175.",
        "python -m unittest discover -s backend/tests -p 'test_*.py' -> passed",
        "python -m unittest discover -s smartgate/tests -p 'test_*.py' -> passed",
        "python -m unittest discover -s tests -p 'test_*.py' -> passed",
        "npm run build -> passed",
    ]
    y = left[1] + 90
    for item in checks:
        draw.ellipse((left[0] + 28, y + 7, left[0] + 44, y + 23), fill="#16a34a")
        y = draw_wrapped(draw, item, (left[0] + 58, y), 700, body_font, "#1e293b", 6) + 10

    draw.text((right[0] + 28, right[1] + 24), "Live API exercise", font=section_font, fill="#111827")
    calls = [
        "POST /users/register -> user_id = CC3C0CD0",
        "GET /zones -> returned correct Chinese labels such as 羽毛球 / 籃球區",
        "POST /zones/occupancy (A=1, B=2) -> zone statuses updated",
        "POST /queue/join -> walk_in = true for low occupancy",
        "POST /session/enter -> session_id = S-A6027D5C",
        "GET /sessions/active -> active countdown data returned",
    ]
    y = right[1] + 90
    for item in calls:
        draw.rectangle((right[0] + 28, y + 4, right[0] + 44, y + 20), fill="#0ea5e9")
        y = draw_wrapped(draw, item, (right[0] + 58, y), 720, body_font, "#1e293b", 6) + 10

    draw.text((70, 950), "The report therefore describes a verified final build, not a paper-only proposal.", font=body_font, fill="#475569")
    img.save(path)
    return path


def make_autonomous_flowchart() -> Path:
    path = ASSET_DIR / "autonomous-flowchart.png"
    img = Image.new("RGB", (1800, 1160), "#f8fafc")
    draw = ImageDraw.Draw(img)
    title_font = font(40, bold=True)
    node_title = font(24, bold=True)
    body_font = font(18)
    small_font = font(16)
    draw.text((70, 40), "Autonomous Queue and Session Control Flow", font=title_font, fill="#0f172a")
    draw.text((70, 92), "This flowchart shows how the final BridgeSpace backend makes decisions during operation.", font=body_font, fill="#475569")

    nodes = {
        "arrive": (680, 140, 1120, 240, "#fff7ed", "#f97316", "Visitor scans at SmartGate", "Kiosk identifies or registers the visitor and sends the zone choice to the backend."),
        "join": (680, 300, 1120, 420, "#eff6ff", "#2563eb", "Queue join request", "Backend checks occupancy, current queue length, and zone policy."),
        "walkin": (1200, 300, 1660, 420, "#ecfdf5", "#16a34a", "Walk-in path", "If occupancy is below the threshold and no queue exists, the visitor can enter immediately."),
        "waiting": (160, 300, 620, 420, "#fefce8", "#ca8a04", "Waiting path", "If the zone is busy, the visitor receives a queue number and waits for the next call."),
        "session": (680, 500, 1120, 620, "#f5f3ff", "#7c3aed", "Session starts", "The backend creates a timed session and returns session_id, duration, and expiry."),
        "count": (160, 500, 620, 620, "#ecfeff", "#0891b2", "SmartCount updates occupancy", "Camera-based counts are pushed back to the backend every cycle."),
        "watcher": (160, 700, 620, 840, "#eff6ff", "#0ea5e9", "OccupancyWatcher decision", "Stable departure detection decides whether a new queue call is safe."),
        "call": (160, 900, 620, 1020, "#fff7ed", "#ea580c", "Auto-call next visitor", "The next waiting visitor is marked called and the dashboard is updated."),
        "timer": (680, 700, 1120, 840, "#fef2f2", "#dc2626", "Session timer checks", "SessionManager enforces warning, expiry, extension, no-show, and overstay rules."),
        "notify": (1200, 700, 1660, 840, "#faf5ff", "#9333ea", "Device + alert actions", "The system can log device actions and send dashboard, Telegram, or Twilio alerts."),
    }

    def node(key):
        x1, y1, x2, y2, fill, outline, title, text = nodes[key]
        draw.rounded_rectangle((x1, y1, x2, y2), radius=28, fill=fill, outline=outline, width=4)
        draw.text((x1 + 24, y1 + 20), title, font=node_title, fill="#111827")
        draw_wrapped(draw, text, (x1 + 24, y1 + 58), x2 - x1 - 48, body_font, "#334155", 6)

    for key in nodes:
        node(key)

    def arrow(start, end, color="#94a3b8", label=None, label_pos=None):
        draw.line((start, end), fill=color, width=7)
        angle = math.atan2(end[1] - start[1], end[0] - start[0])
        arrow_size = 16
        p1 = (end[0] - arrow_size * math.cos(angle - math.pi / 6), end[1] - arrow_size * math.sin(angle - math.pi / 6))
        p2 = (end[0] - arrow_size * math.cos(angle + math.pi / 6), end[1] - arrow_size * math.sin(angle + math.pi / 6))
        draw.polygon([end, p1, p2], fill=color)
        if label and label_pos:
            draw.text(label_pos, label, font=small_font, fill="#475569")

    arrow((900, 240), (900, 300), label="send zone choice", label_pos=(930, 255))
    arrow((680, 360), (620, 360), label="busy", label_pos=(630, 330))
    arrow((1120, 360), (1200, 360), label="walk-in allowed", label_pos=(1160, 330))
    arrow((1430, 420), (1000, 500), label="enter zone", label_pos=(1180, 450))
    arrow((900, 420), (900, 500))
    arrow((390, 420), (390, 500))
    arrow((390, 620), (390, 700), label="new count", label_pos=(420, 640))
    arrow((620, 770), (680, 770), label="end queue wait", label_pos=(620, 735))
    arrow((900, 620), (900, 700), label="check every cycle", label_pos=(930, 640))
    arrow((1120, 770), (1200, 770), label="warning / expiry", label_pos=(1130, 735))
    arrow((390, 840), (390, 900), label="slot available", label_pos=(420, 860))
    arrow((620, 960), (680, 560), label="called user enters", label_pos=(640, 860))
    arrow((1200, 770), (1120, 770), color="#c084fc")

    draw.text((70, 1080), "This flow connects the presentation story to the code: kiosk requests, occupancy updates, queue advancement, session timing, and alerting are all implemented in the final repository.", font=body_font, fill="#475569")
    img.save(path)
    return path


def make_live_evidence_panel() -> Path:
    path = ASSET_DIR / "live-evidence-panel.png"
    dashboard = Image.open(ROOT / "tmp-live-assets" / "dashboard-live-wait.png").convert("RGB")
    backend_docs = Image.open(ROOT / "tmp-live-assets" / "backend-docs.png").convert("RGB")
    sessions = Image.open(ROOT / "tmp-live-assets" / "api-sessions.png").convert("RGB")
    queue = Image.open(ROOT / "tmp-live-assets" / "api-queue.png").convert("RGB")

    # Crop to the informative parts.
    dashboard = dashboard.crop((0, 90, 1220, 690))
    backend_docs = backend_docs.crop((60, 40, 1510, 840))
    sessions = sessions.crop((0, 0, 1500, 220))
    queue = queue.crop((0, 0, 1500, 180))

    canvas = Image.new("RGB", (1800, 1280), "#ffffff")
    draw = ImageDraw.Draw(canvas)
    title_font = font(40, bold=True)
    caption_font = font(18)
    section_font = font(24, bold=True)
    draw.text((60, 40), "Actual Running System Evidence", font=title_font, fill="#0f172a")
    draw.text((60, 92), "Captured from the final local environment on 12 April 2026.", font=caption_font, fill="#475569")

    def paste_fit(image, box):
        x1, y1, x2, y2 = box
        copy = image.copy()
        copy.thumbnail((x2 - x1, y2 - y1))
        px = x1 + ((x2 - x1) - copy.width) // 2
        py = y1 + ((y2 - y1) - copy.height) // 2
        canvas.paste(copy, (px, py))
        draw.rounded_rectangle(box, radius=20, outline="#cbd5e1", width=3)

    left_box = (60, 150, 980, 760)
    right_top = (1030, 150, 1740, 720)
    right_mid = (1030, 760, 1740, 980)
    right_bottom = (1030, 1000, 1740, 1210)

    paste_fit(dashboard, left_box)
    paste_fit(backend_docs, right_top)
    paste_fit(sessions, right_mid)
    paste_fit(queue, right_bottom)

    draw.text((80, 170), "Live frontend dashboard", font=section_font, fill="#111827")
    draw.text((1050, 170), "Live backend Swagger /docs", font=section_font, fill="#111827")
    draw.text((1050, 780), "GET /sessions/active response", font=section_font, fill="#111827")
    draw.text((1050, 1020), "GET /queue/A response", font=section_font, fill="#111827")

    draw.text((80, 720), "The dashboard screenshot was captured from the running Vite frontend and shows real occupancy values fed by the active backend.", font=caption_font, fill="#475569")
    draw.text((1050, 680), "The FastAPI docs confirm the full v2.0 endpoint set used in the demo: registration, queue, occupancy, session, device, and alert routes.", font=caption_font, fill="#475569")
    draw.text((1050, 948), "This actual response shows active timed sessions being returned by the running backend.", font=caption_font, fill="#475569")
    draw.text((1050, 1188), "This actual response shows queue entries waiting in Zone A during the demo setup.", font=caption_font, fill="#475569")

    canvas.save(path)
    return path


def make_styles(document: Document) -> None:
    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    if "BridgeSpace Caption" not in styles:
        style = styles.add_style("BridgeSpace Caption", WD_STYLE_TYPE.PARAGRAPH)
        style.font.name = "Calibri"
        style.font.size = Pt(9.5)
        style.font.italic = True
        style.font.color.rgb = RGBColor(90, 90, 90)


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def page_number(paragraph):
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_end)


def add_heading(document: Document, text: str, level: int = 1):
    paragraph = document.add_paragraph()
    paragraph.style = document.styles[f"Heading {level}"]
    paragraph.paragraph_format.space_before = Pt(10)
    paragraph.paragraph_format.space_after = Pt(4)
    run = paragraph.add_run(text)
    if level == 1:
        run.font.color.rgb = RGBColor(15, 23, 42)
    return paragraph


def add_caption(document: Document, text: str):
    p = document.add_paragraph(style="BridgeSpace Caption")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(text)


def add_bullet_list(document: Document, items: list[str]):
    for item in items:
        p = document.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(0)
        p.add_run(item)


def add_table(document: Document, headers: list[str], rows: list[list[str]]):
    table = document.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for idx, header in enumerate(headers):
        hdr_cells[idx].text = header
        hdr_cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_shading(hdr_cells[idx], "1F2937")
        for paragraph in hdr_cells[idx].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = value
            cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    return table


def add_image(document: Document, image_path: Path, width: float):
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(image_path), width=Inches(width))


def build_report():
    architecture = make_architecture_diagram()
    operation = make_operation_diagram()
    validation = make_validation_panel()
    control_flow = make_autonomous_flowchart()
    live_evidence = make_live_evidence_panel()

    site_photo_1 = ROOT / "實際地點圖片" / "WhatsApp Image 2026-03-17 at 4.44.08 PM.jpeg"
    site_photo_2 = ROOT / "實際地點圖片" / "WhatsApp Image 2026-03-17 at 4.44.08 PM (1).jpeg"
    concept_image = ROOT / "BridgeSpace_AI_System.jpg"
    poster_image = ROOT / "BridgeSpace_Final_Poster.png"

    document = Document()
    make_styles(document)

    section = document.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("BridgeSpace Final Report")
    run.font.size = Pt(24)
    run.font.bold = True
    run.font.color.rgb = RGBColor(15, 23, 42)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("Updated final submission based on the completed v2.0 build and final practical verification").font.size = Pt(12)

    meta = document.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run("COM1002 Cyber Technology and Society | Group 5 | The Hang Seng University of Hong Kong\nFinal exhibition preparation update | 12 April 2026").font.size = Pt(11)

    add_image(document, concept_image, 6.6)
    add_caption(document, "Figure 1. BridgeSpace concept visual showing the underpass sports hub and AI-assisted entry experience.")

    intro = document.add_paragraph()
    intro.paragraph_format.space_before = Pt(10)
    intro.paragraph_format.space_after = Pt(0)
    intro.add_run(
        "BridgeSpace is a service-learning proposal and working prototype that transforms idle space beneath Sha Lek Highway in Sha Tin into an AI-powered community sports hub. "
        "This updated report replaces the previous version by reflecting the final codebase, the repaired v2.0 workflows, and the practical validation completed on 12 April 2026."
    )

    document.add_page_break()

    add_heading(document, "1. Executive Summary", 1)
    document.add_paragraph(
        "BridgeSpace addresses a clear community problem in Sha Tin: a large residential population competes for limited sports space, while conventional booking systems are vulnerable to scalping and unfair access. "
        "The project proposes a walk-in-only, AI-assisted sports hub that uses underpass space more effectively and shifts access from speculative online booking to physical presence."
    )
    document.add_paragraph(
        "The final v2.0 build combines four main layers. SmartGate provides kiosk-based registration, recognition, queue entry, and session extension. SmartCount provides occupancy input from computer vision. "
        "A FastAPI backend coordinates users, zones, queues, sessions, device logs, and alerts. A React dashboard presents the facility state live through WebSocket updates. "
        "Together, these parts form a closed operational loop in which occupancy changes, session rules, queue progression, and alerting can happen without continuous staff intervention."
    )
    document.add_paragraph(
        "This updated report also includes practical evidence. The final local run confirmed backend startup, frontend startup, registration, correct zone label output, occupancy updates, walk-in admission, session creation, and active session retrieval. "
        "The report therefore documents a verified prototype rather than a purely conceptual design."
    )

    add_heading(document, "2. Community Need and Site Opportunity", 1)
    document.add_paragraph(
        "Our original community-needs work identified dissatisfaction with sports availability in Sha Tin and highlighted the fairness problem caused by court-booking scalping. "
        "The underpass beneath Sha Lek Highway offers a compelling alternative because it is already sheltered, linearly organised by structural bays, and physically close to residential foot traffic."
    )
    add_image(document, site_photo_1, 6.5)
    add_caption(document, "Figure 2. Existing underpass condition at the Sha Tin site selected for BridgeSpace.")
    document.add_paragraph(
        "The site is suitable for modular conversion. Each structural bay can host a sports or support function, while the existing bridge deck provides weather protection. "
        "This lowers the conceptual cost of enclosure and supports a phased implementation model."
    )
    add_image(document, site_photo_2, 6.5)
    add_caption(document, "Figure 3. A second site photo illustrating the spatial depth and repetitive bay structure that supports modular planning.")

    add_heading(document, "3. Final Solution Overview", 1)
    document.add_paragraph(
        "The final BridgeSpace concept combines a physical underpass hub with an autonomous digital management system. "
        "Visitors arrive on site, identify themselves at the kiosk, choose a zone, and either walk in directly or join a queue. "
        "Once in a zone, timed sessions are enforced through the backend and reflected on the live dashboard. "
        "When occupancy changes, the system can automatically advance the queue and notify waiting users."
    )
    add_image(document, poster_image, 4.2)
    add_caption(document, "Figure 4. Final poster composition summarising the BridgeSpace proposition, visual identity, and user-facing concept.")

    add_heading(document, "4. Final System Architecture", 1)
    document.add_paragraph(
        "The architecture below reflects the implemented state of the current repository rather than the older draft architecture. "
        "Compared with the previous report, the final build includes repaired kiosk-to-backend contracts, normalised zone labels, cleaned dashboard copy, and regression tests that protect these fixes."
    )
    add_image(document, architecture, 6.9)
    add_caption(document, "Figure 5. Final digital architecture of BridgeSpace v2.0 based on the code running in the completed repository.")
    add_table(
        document,
        ["Layer", "Current responsibility"],
        [
            ["SmartGate kiosk", "User registration, local face-signature match, queue join, walk-in flow, session extension"],
            ["SmartCount", "Zone occupancy updates via camera-based counting"],
            ["FastAPI backend", "Users, zones, queue, sessions, alerts, device log, WebSocket updates"],
            ["SessionManager", "Session duration, warning, expiry, extension, and no-show rules"],
            ["OccupancyWatcher", "Departure confirmation and auto queue advancement"],
            ["Dashboard", "Live occupancy, session, queue, device, and alert visualisation"],
        ],
    )

    add_heading(document, "5. Practical Operation and Demonstration", 1)
    document.add_paragraph(
        "The project is now documented not only as a design but also as an executable operator workflow. "
        "The final runbook starts the system in four parts: backend API, SmartCount, SmartGate, and the React dashboard. "
        "This order ensures the dashboard always connects to a stable backend and that occupancy updates can affect queue and session logic during a live demonstration."
    )
    add_image(document, operation, 6.9)
    add_caption(document, "Figure 6. Verified operation flow used to explain and demonstrate the final prototype.")
    add_bullet_list(document, [
        "Backend startup: `python -m uvicorn main:app --host 127.0.0.1 --port 8012` and `--port 8000` for dashboard compatibility.",
        "Frontend startup: `npm run dev -- --host 127.0.0.1 --port 4175`.",
        "Verified flow: register visitor -> retrieve zones -> update occupancy -> join queue / walk in -> enter session -> retrieve active sessions.",
        "Verified response examples: `walk_in = true`, session created successfully, correct Chinese zone labels returned from `/zones`.",
    ])
    add_image(document, control_flow, 6.9)
    add_caption(document, "Figure 7. Detailed control flow for queue advancement, session timing, and backend decision-making in the final build.")

    add_heading(document, "6. Actual Running System Images", 1)
    document.add_paragraph(
        "The following figure is included to address the most important reporting requirement: the report must show the intelligent system actually running. "
        "Instead of relying only on concept art, this page combines real screenshots from the local final environment. "
        "It includes the live frontend dashboard, the running backend Swagger documentation, and actual API responses for active sessions and queue state."
    )
    add_image(document, live_evidence, 6.9)
    add_caption(document, "Figure 8. Real screenshots captured from the final local run, showing frontend output and backend responses.")
    document.add_paragraph(
        "These screenshots matter because they link the report directly to practical operation. "
        "The dashboard image demonstrates a rendered frontend connected to real occupancy data, while the backend screenshots show that the service was not mocked: it exposed the documented endpoints and returned active session and queue data during the same run."
    )
    add_image(document, validation, 6.9)
    add_caption(document, "Figure 9. Practical verification summary compiled from the same final run used for the screenshots above.")

    add_heading(document, "7. Demonstration Script and Explanation", 1)
    document.add_paragraph(
        "A clean live demonstration should follow a clear script so the audience can understand what each subsystem is doing. "
        "The sequence below is suitable for an exhibition or presentation setting."
    )
    add_bullet_list(document, [
        "Step 1. Open the dashboard and explain that it reflects live zone occupancy, queue state, sessions, devices, and alerts.",
        "Step 2. Show the backend `/docs` page to establish that the system is running as a real API service rather than a slide-only concept.",
        "Step 3. Register or identify a visitor through SmartGate, then select a zone and explain why the backend either allows walk-in or returns a queue number.",
        "Step 4. Trigger occupancy updates and explain how SmartCount affects the queue and the displayed zone bars.",
        "Step 5. Show `/sessions/active` and explain how the backend stores and returns timed session information.",
        "Step 6. Show `/queue/A` or another queue endpoint to explain how waiting visitors are represented and how the next call is decided.",
        "Step 7. Summarise the autonomous loop: sensing, decision, action, and notification.",
    ])

    add_heading(document, "8. What Was Fixed in the Final Completion Pass", 1)
    document.add_paragraph(
        "A final review identified issues that would have undermined the exhibition build if they had remained unresolved. "
        "These were addressed before this report was rewritten."
    )
    add_table(
        document,
        ["Issue", "Final correction"],
        [
            ["SmartGate dependency drift", "Replaced the brittle hard dependency with a local face-signature fallback suitable for demo use."],
            ["Broken zone labels", "Normalised zone data so the backend returns correct Chinese names instead of mojibake."],
            ["Frontend user-facing copy corruption", "Cleaned header, footer, queue-call text, and related dashboard labels."],
            ["Missing protection against regressions", "Added backend, SmartGate, and copy-focused test coverage."],
        ],
    )

    add_heading(document, "9. Verification and Testing", 1)
    document.add_paragraph(
        "Before this report was finalised, the completed repository was checked through both automated and live validation. "
        "The key successful checks are listed below."
    )
    add_bullet_list(document, [
        "Python unit tests passed for backend zone catalog handling.",
        "Python unit tests passed for SmartGate face-matching fallback behaviour.",
        "Python unit tests passed for user-facing frontend copy checks.",
        "Frontend production build passed with `npm run build`.",
        "Live API exercise confirmed successful register, queue, session entry, and active session retrieval flows.",
    ])
    document.add_paragraph(
        "This combination matters because a service-learning system like BridgeSpace must be explainable in presentation form and also defensible as a functioning prototype. "
        "The final state now supports both goals."
    )

    add_heading(document, "10. Limitations and Future Work", 1)
    document.add_paragraph(
        "The current SmartGate fallback is appropriate for a demonstration or course project, but it is not production-grade biometric security. "
        "A real deployment would require stronger identity controls, improved privacy governance, and more robust hardware integration."
    )
    add_bullet_list(document, [
        "Replace demo-grade local face matching with a production-ready identity approach.",
        "Connect SmartControl to real hardware rather than simulated device states.",
        "Deploy backend and dashboard to stable public infrastructure for field trials.",
        "Expand occupancy calibration and multi-zone camera coverage for real underpass use.",
    ])

    add_heading(document, "11. Conclusion", 1)
    document.add_paragraph(
        "BridgeSpace now stands as a completed v2.0 exhibition build with a verified operational loop, supporting documents, and corrected implementation details. "
        "It demonstrates how underused urban infrastructure can be transformed into a fair-access community sports service through a combination of spatial design, software, and automation."
    )
    document.add_paragraph(
        "More importantly, the project now communicates that story clearly. "
        "The previous report described the proposal; this updated report describes the proposal, the implemented system, the final repair work, and the practical evidence that the system can run."
    )

    footer = document.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run("BridgeSpace Final Report Updated | Page ")
    page_number(footer)

    document.save(OUTPUT)


if __name__ == "__main__":
    build_report()
    print(json.dumps({"output": str(OUTPUT), "assets": str(ASSET_DIR)}, ensure_ascii=False))
